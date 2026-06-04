from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

# ===== PORTADA =====
for _ in range(6):
    doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('Sistema de Gestion y Alquiler de Productos')
run.bold = True
run.font.size = Pt(24)
run.font.color.rgb = RGBColor(0, 51, 102)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub.add_run('Documento de Requisitos y Especificacion Tecnica')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(100, 100, 100)

doc.add_paragraph()
info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = info.add_run('Backend API REST - Flask + PostgreSQL + MongoDB')
run.font.size = Pt(12)

doc.add_page_break()

# ===== INDICE =====
doc.add_heading('Indice', level=1)
items = [
    '1. Resumen del Proyecto',
    '2. Requisitos Funcionales',
    '   2.1. Modulo de Autenticacion',
    '   2.2. Modulo de Usuarios',
    '   2.3. Modulo de Productos y Categorias',
    '   2.4. Modulo de Ubicaciones',
    '   2.5. Modulo de Roles',
    '   2.6. Modulo de Favoritos',
    '   2.7. Modulo de Reservas',
    '   2.8. Modulo de Pagos',
    '   2.9. Modulo de Facturas',
    '   2.10. Modulo de Resenas',
    '   2.11. Modulo de Puntos',
    '3. Requisitos No Funcionales',
    '4. Arquitectura del Sistema',
    '5. Modelo de Datos',
    '   5.1. Diagrama Entidad-Relacion (PostgreSQL)',
    '   5.2. Colecciones MongoDB',
    '6. Endpoints de la API',
    '7. Seguridad',
    '8. Tecnologias Utilizadas',
    '9. Dependencias'
]
for item in items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)

doc.add_page_break()

# ===== 1. RESUMEN =====
doc.add_heading('1. Resumen del Proyecto', level=1)
doc.add_paragraph(
    'El Sistema de Gestion y Alquiler de Productos es una API REST disenada para '
    'permitir la publicacion, busqueda y reserva de productos en alquiler. '
    'Implementa una arquitectura de persistencia poliglota combinando PostgreSQL '
    'para datos transaccionales (usuarios, reservas, pagos) y MongoDB para el '
    'catalogo de productos y categorias, aprovechando la flexibilidad de documentos '
    'para almacenar especificaciones tecnicas heterogeneas.'
)
doc.add_paragraph(
    'La API expone 49 endpoints organizados en 11 modulos funcionales, '
    'con autenticacion JWT, paginacion, validacion de datos y manejo de errores estandarizado.'
)

# ===== 2. REQUISITOS FUNCIONALES =====
doc.add_heading('2. Requisitos Funcionales', level=1)

modulos = [
    ('2.1. Modulo de Autenticacion', 'auth_routes.py', [
        'RF-01: El sistema debe permitir el registro de usuarios con nombre, email, contrasena y telefono.',
        'RF-02: El sistema debe hashear la contrasena antes de almacenarla (werkzeug.security).',
        'RF-03: El sistema debe rechazar registros con email duplicado.',
        'RF-04: El sistema debe asignar automaticamente el rol "cliente" al nuevo usuario.',
        'RF-05: El sistema debe devolver un token JWT al registrar o iniciar sesion.',
        'RF-06: El sistema debe validar credenciales (email + contrasena) en el inicio de sesion.',
        'RF-07: El sistema debe exponer un endpoint protegido para obtener el perfil del usuario autenticado.',
        'RF-08: Los tokens JWT deben expirar a las 24 horas.',
    ]),
    ('2.2. Modulo de Usuarios', 'user_routes.py', [
        'RF-09: El sistema debe listar usuarios con paginacion (page, per_page).',
        'RF-10: El sistema debe permitir obtener un usuario por su ID.',
        'RF-11: El sistema debe permitir crear usuarios (solo admins con token).',
        'RF-12: El sistema debe permitir actualizar datos del usuario (nombre, email, telefono, contrasena).',
        'RF-13: El sistema debe permitir eliminar usuarios con borrado en cascada.',
    ]),
    ('2.3. Modulo de Productos y Categorias', 'product_routes.py + mongo_models.py', [
        'RF-14: El sistema debe permitir crear categorias en MongoDB con nombre y descripcion.',
        'RF-15: El sistema debe rechazar categorias con nombre duplicado.',
        'RF-16: El sistema debe listar categorias paginadas desde MongoDB.',
        'RF-17: El sistema debe permitir crear productos vinculados a una categoria (ObjectId).',
        'RF-18: El sistema debe listar productos paginados desde MongoDB.',
        'RF-19: El sistema debe filtrar productos por categoria.',
        'RF-20: El sistema debe obtener un producto individual por su ID de MongoDB.',
        'RF-21: Al crear un producto, el usuario propietario debe recibir el rol "rentador".',
    ]),
    ('2.4. Modulo de Ubicaciones', 'location_routes.py', [
        'RF-22: El sistema debe crear ubicaciones asociadas a un usuario.',
        'RF-23: El sistema debe validar rangos de latitud (-90 a 90) y longitud (-180 a 180).',
        'RF-24: El sistema debe listar ubicaciones paginadas.',
        'RF-25: El sistema debe filtrar ubicaciones por usuario.',
        'RF-26: El sistema debe eliminar ubicaciones.',
    ]),
    ('2.5. Modulo de Roles', 'user_role_routes.py', [
        'RF-27: El sistema debe asignar roles a usuarios.',
        'RF-28: El sistema debe rechazar asignaciones duplicadas del mismo rol al mismo usuario.',
        'RF-29: El sistema debe listar asignaciones de roles paginadas.',
        'RF-30: El sistema debe listar roles de un usuario especifico incluyendo el nombre del rol.',
        'RF-31: El sistema debe eliminar asignaciones de roles.',
        'RF-32: Los roles predefinidos son "cliente" y "rentador".',
    ]),
    ('2.6. Modulo de Favoritos', 'favorite_routes.py', [
        'RF-33: El sistema debe agregar productos a favoritos de un usuario.',
        'RF-34: El sistema debe validar que el usuario exista en PostgreSQL.',
        'RF-35: El sistema debe validar que el producto exista en MongoDB.',
        'RF-36: El sistema debe rechazar favoritos duplicados.',
        'RF-37: El sistema debe listar favoritos paginados.',
        'RF-38: El sistema debe filtrar favoritos por usuario.',
        'RF-39: El sistema debe eliminar favoritos.',
    ]),
    ('2.7. Modulo de Reservas', 'reservation_routes.py', [
        'RF-40: El sistema debe crear reservas vinculando un usuario (PostgreSQL) con un producto (MongoDB).',
        'RF-41: El sistema debe validar fechas: la fecha de inicio debe ser anterior a la fecha de fin.',
        'RF-42: El sistema debe permitir actualizar el estado, fechas, precio y ubicacion de una reserva.',
        'RF-43: El sistema debe listar reservas paginadas.',
        'RF-44: El sistema debe filtrar reservas por usuario.',
        'RF-45: El sistema debe eliminar reservas.',
    ]),
    ('2.8. Modulo de Pagos', 'payment_routes.py', [
        'RF-46: El sistema debe registrar pagos asociados a una reserva.',
        'RF-47: El sistema debe generar un codigo de transaccion unico (UUID) si no se proporciona.',
        'RF-48: El sistema debe listar pagos paginados.',
        'RF-49: El sistema debe eliminar pagos.',
    ]),
    ('2.9. Modulo de Facturas', 'invoice_routes.py', [
        'RF-50: El sistema debe generar facturas asociadas a un pago.',
        'RF-51: El sistema debe validar que el total sea mayor a cero.',
        'RF-52: El sistema debe generar un codigo de factura unico si no se proporciona.',
        'RF-53: El sistema debe listar facturas paginadas.',
        'RF-54: El sistema debe eliminar facturas.',
    ]),
    ('2.10. Modulo de Resenas', 'review_routes.py', [
        'RF-55: El sistema debe crear resenas con calificacion (1-5) y comentario.',
        'RF-56: El sistema debe validar que la calificacion este entre 1 y 5.',
        'RF-57: El sistema debe validar que la reserva, el reviewer y el usuario evaluado existan.',
        'RF-58: El sistema debe listar resenas paginadas.',
        'RF-59: El sistema debe filtrar resenas por usuario evaluado.',
        'RF-60: El sistema debe eliminar resenas.',
    ]),
    ('2.11. Modulo de Puntos', 'point_routes.py', [
        'RF-61: El sistema debe registrar puntos de fidelidad para usuarios.',
        'RF-62: El sistema debe listar registros de puntos paginados.',
        'RF-63: El sistema debe mostrar el total acumulado de puntos por usuario.',
        'RF-64: El sistema debe eliminar registros de puntos.',
    ]),
]

for titulo, archivo, requisitos in modulos:
    doc.add_heading(titulo, level=2)
    p = doc.add_paragraph()
    run = p.add_run(f'Archivo: {archivo}')
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(100, 100, 100)
    for req in requisitos:
        doc.add_paragraph(req, style='List Bullet')

doc.add_page_break()

# ===== 3. REQUISITOS NO FUNCIONALES =====
doc.add_heading('3. Requisitos No Funcionales', level=1)
no_func = [
    ('RNF-01', 'Arquitectura REST', 'La API debe seguir los principios REST con metodos HTTP estandar (GET, POST, PUT, DELETE).'),
    ('RNF-02', 'Persistencia Poliglota', 'Usar PostgreSQL para datos transaccionales y MongoDB para catalogos flexibles.'),
    ('RNF-03', 'Autenticacion JWT', 'Todas las operaciones de escritura deben requerir autenticacion mediante tokens JWT.'),
    ('RNF-04', 'Paginacion', 'Todos los endpoints de listado deben soportar paginacion via parametros page y per_page.'),
    ('RNF-05', 'Validacion de Datos', 'Todos los endpoints deben validar campos requeridos y tipos de datos.'),
    ('RNF-06', 'Manejo de Errores', 'Los errores deben devolverse en formato JSON estandar con codigos HTTP apropiados.'),
    ('RNF-07', 'Seguridad de Contrasenas', 'Las contrasenas deben almacenarse hasheadas con werkzeug.security.'),
    ('RNF-08', 'Escalabilidad', 'La arquitectura debe permitir escalar horizontalmente separando responsabilidades.'),
    ('RNF-09', 'Portabilidad', 'El sistema debe ejecutarse en cualquier entorno con Python 3.10+ y acceso a las bases de datos.'),
    ('RNF-10', 'Modularidad', 'Cada modulo funcional debe implementarse como un Blueprint independiente de Flask.'),
]
for cod, nombre, desc in no_func:
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Light Grid Accent 1'
    table.cell(0, 0).text = cod
    table.cell(0, 1).text = nombre
    table.cell(0, 2).text = desc
    table.cell(0, 0).width = Inches(1)
    table.cell(0, 1).width = Inches(1.5)
    # bold the code
    for run in table.cell(0, 0).paragraphs[0].runs:
        run.bold = True
    doc.add_paragraph()

# ===== 4. ARQUITECTURA =====
doc.add_page_break()
doc.add_heading('4. Arquitectura del Sistema', level=1)
doc.add_paragraph(
    'El sistema sigue una arquitectura de 3 capas implementada con Flask:'
)

capas = [
    ('Capa de Presentacion (API)', 
     'Blueprints de Flask que exponen endpoints REST. Cada modulo (auth, users, products, etc.) '
     'tiene su propio blueprint registrado con un prefijo de URL.'),
    ('Capa de Negocio (Logica)', 
     'Implementada directamente en las rutas y modelos. Incluye validaciones, '
     'asignacion de roles, calculo de totales y logica de negocios especifica.'),
    ('Capa de Datos (Persistencia)', 
     'SQLAlchemy para PostgreSQL (10 tablas) y PyMongo para MongoDB (2 colecciones principales). '
     'La conexion a MongoDB se maneja mediante una variable global en la fabrica de la aplicacion.'),
]

for titulo, desc in capas:
    p = doc.add_paragraph()
    run = p.add_run(titulo + ': ')
    run.bold = True
    p.add_run(desc)

doc.add_heading('4.1. Flujo de Autenticacion', level=2)
doc.add_paragraph(
    '1. El cliente envia POST /api/auth/register con full_name, email y password.\n'
    '2. El servidor hashea la password con generate_password_hash().\n'
    '3. Crea el registro en PostgreSQL (tabla users).\n'
    '4. Asigna rol "cliente" en user_roles.\n'
    '5. Genera un token JWT con create_access_token() y lo devuelve.\n'
    '6. Para endpoints protegidos, el cliente envia el token en header Authorization: Bearer <token>.\n'
    '7. El decorador @jwt_required() verifica el token antes de ejecutar el handler.'
)

doc.add_heading('4.2. Flujo de Creacion de Producto con Cambio de Rol', level=2)
doc.add_paragraph(
    '1. Usuario autenticado envia POST /api/products con owner_id, name_prod, category_id, price.\n'
    '2. Se inserta el producto en MongoDB (coleccion products).\n'
    '3. Se verifica si el usuario ya tiene rol "rentador".\n'
    '4. Si no lo tiene, se le asigna (conservando el rol "cliente").\n'
    '5. El usuario ahora puede publicar productos como rentador.'
)

# ===== 5. MODELO DE DATOS =====
doc.add_page_break()
doc.add_heading('5. Modelo de Datos', level=1)

doc.add_heading('5.1. Diagrama Entidad-Relacion (PostgreSQL)', level=2)

# Tabla de modelos
tablas = [
    ('users', 'id, full_name, email (UNIQUE), phone, password_hash, reputation_score, created_at'),
    ('roles', 'id, role_name (UNIQUE)'),
    ('user_roles', 'id, user_id (FK->users), role_id (FK->roles), assigned_at'),
    ('locations', 'id, user_id (FK->users), location_name, address, latitude, longitude, city, reference, is_meeting_point, created_at'),
    ('reservations', 'id, renter_user_id (FK->users), mongo_product_id, pickup_location_id (FK->locations), start_date, end_date, total_price, status, created_at'),
    ('payments', 'id, reservation_id (FK->reservations), amount, payment_date, payment_status, transaction_code (UNIQUE)'),
    ('invoices', 'id, payment_id (FK->payments), subtotal, platform_fee, owner_earnings, taxes, total, invoice_code (UNIQUE), generated_at'),
    ('reviews', 'id, reservation_id (FK->reservations), reviewer_id (FK->users), reviewed_user_id (FK->users), rating, comment, mongo_review_id, created_at'),
    ('points', 'id, user_id (FK->users), points_earned, reason, created_at'),
    ('favorites', 'id, user_id (FK->users), mongo_product_id, created_at'),
]

table = doc.add_table(rows=len(tablas)+1, cols=2)
table.style = 'Light Grid Accent 1'
table.cell(0, 0).text = 'Tabla'
table.cell(0, 1).text = 'Columnas'
for run in table.cell(0, 0).paragraphs[0].runs:
    run.bold = True
for run in table.cell(0, 1).paragraphs[0].runs:
    run.bold = True

for i, (t, c) in enumerate(tablas, 1):
    table.cell(i, 0).text = t
    table.cell(i, 1).text = c

doc.add_paragraph()
doc.add_paragraph('Relaciones entre tablas:')
relaciones = [
    'users 1:N locations (un usuario tiene muchas ubicaciones)',
    'users 1:N reservations (un usuario renta muchos productos)',
    'users 1:N reviews como reviewer',
    'users 1:N reviews como reviewed_user',
    'users 1:N points (un usuario acumula puntos)',
    'users 1:N favorites (un usuario tiene favoritos)',
    'users N:M roles via user_roles',
    'reservations 1:1 payments (una reserva tiene un pago)',
    'payments 1:1 invoices (un pago genera una factura)',
    'locations 1:N reservations (una ubicacion puede ser punto de entrega)',
]
for r in relaciones:
    doc.add_paragraph(r, style='List Bullet')

doc.add_heading('5.2. Colecciones MongoDB', level=2)
mongo_cols = [
    ('categories', 'Documento flexible con _id (ObjectId), name_cat (unicode indexado), description'),
    ('products', 'Documento con _id (ObjectId), owner_id (ref PostgreSQL), name_prod, description, category_id (ObjectId ref categories), price, details, status'),
]

table = doc.add_table(rows=len(mongo_cols)+1, cols=2)
table.style = 'Light Grid Accent 1'
table.cell(0, 0).text = 'Coleccion'
table.cell(0, 1).text = 'Estructura'
for run in table.cell(0, 0).paragraphs[0].runs:
    run.bold = True
for run in table.cell(0, 1).paragraphs[0].runs:
    run.bold = True
for i, (t, c) in enumerate(mongo_cols, 1):
    table.cell(i, 0).text = t
    table.cell(i, 1).text = c

doc.add_page_break()

# ===== 6. ENDPOINTS =====
doc.add_heading('6. Endpoints de la API', level=1)
doc.add_paragraph('Total: 49 endpoints organizados en 11 blueprints.')

endpoints_data = [
    ('Auth', [
        ('POST', '/api/auth/register', 'Registrar usuario (publico)', 'No'),
        ('POST', '/api/auth/login', 'Iniciar sesion (publico)', 'No'),
        ('GET', '/api/auth/me', 'Obtener perfil actual', 'Si'),
    ]),
    ('Usuarios', [
        ('GET', '/api/users', 'Listar usuarios (paginado)', 'No'),
        ('GET', '/api/users/<id>', 'Obtener usuario por ID', 'No'),
        ('POST', '/api/users', 'Crear usuario', 'Si'),
        ('PUT', '/api/users/<id>', 'Actualizar usuario', 'Si'),
        ('DELETE', '/api/users/<id>', 'Eliminar usuario', 'Si'),
    ]),
    ('Categorias', [
        ('GET', '/api/categories', 'Listar categorias (paginado)', 'No'),
        ('POST', '/api/categories', 'Crear categoria', 'Si'),
    ]),
    ('Productos', [
        ('GET', '/api/products', 'Listar productos (paginado, filtro por categoria)', 'No'),
        ('GET', '/api/products/<id>', 'Obtener producto por ID', 'No'),
        ('POST', '/api/products', 'Crear producto', 'Si'),
    ]),
    ('Ubicaciones', [
        ('GET', '/api/locations', 'Listar ubicaciones (paginado)', 'No'),
        ('GET', '/api/locations/user/<id>', 'Ubicaciones por usuario', 'No'),
        ('POST', '/api/locations', 'Crear ubicacion', 'Si'),
        ('DELETE', '/api/locations/<id>', 'Eliminar ubicacion', 'Si'),
    ]),
    ('Roles', [
        ('GET', '/api/user-roles', 'Listar roles asignados (paginado)', 'No'),
        ('GET', '/api/user-roles/user/<id>', 'Roles por usuario', 'No'),
        ('POST', '/api/user-roles', 'Asignar rol', 'Si'),
        ('DELETE', '/api/user-roles/<id>', 'Eliminar rol asignado', 'Si'),
    ]),
    ('Favoritos', [
        ('GET', '/api/favorites', 'Listar favoritos (paginado)', 'No'),
        ('GET', '/api/favorites/user/<id>', 'Favoritos por usuario', 'No'),
        ('POST', '/api/favorites', 'Agregar favorito', 'Si'),
        ('DELETE', '/api/favorites/<id>', 'Eliminar favorito', 'Si'),
    ]),
    ('Reservas', [
        ('GET', '/api/reservations', 'Listar reservas (paginado)', 'No'),
        ('GET', '/api/reservations/<id>', 'Obtener reserva', 'No'),
        ('GET', '/api/reservations/user/<id>', 'Reservas por usuario', 'No'),
        ('POST', '/api/reservations', 'Crear reserva', 'Si'),
        ('PUT', '/api/reservations/<id>', 'Actualizar reserva', 'Si'),
        ('DELETE', '/api/reservations/<id>', 'Eliminar reserva', 'Si'),
    ]),
    ('Pagos', [
        ('GET', '/api/payments', 'Listar pagos (paginado)', 'No'),
        ('GET', '/api/payments/<id>', 'Obtener pago', 'No'),
        ('POST', '/api/payments', 'Registrar pago', 'Si'),
        ('DELETE', '/api/payments/<id>', 'Eliminar pago', 'Si'),
    ]),
    ('Facturas', [
        ('GET', '/api/invoices', 'Listar facturas (paginado)', 'No'),
        ('GET', '/api/invoices/<id>', 'Obtener factura', 'No'),
        ('POST', '/api/invoices', 'Crear factura', 'Si'),
        ('DELETE', '/api/invoices/<id>', 'Eliminar factura', 'Si'),
    ]),
    ('Resenas', [
        ('GET', '/api/reviews', 'Listar resenas (paginado)', 'No'),
        ('GET', '/api/reviews/user/<id>', 'Resenas por usuario', 'No'),
        ('POST', '/api/reviews', 'Crear resena', 'Si'),
        ('DELETE', '/api/reviews/<id>', 'Eliminar resena', 'Si'),
    ]),
    ('Puntos', [
        ('GET', '/api/points', 'Listar puntos (paginado)', 'No'),
        ('GET', '/api/points/user/<id>', 'Puntos por usuario (con total)', 'No'),
        ('POST', '/api/points', 'Agregar puntos', 'Si'),
        ('DELETE', '/api/points/<id>', 'Eliminar puntos', 'Si'),
    ]),
    ('Utilidad', [
        ('GET', '/api/testmongo', 'Test conexion MongoDB', 'No'),
    ]),
]

for modulo, endpoints in endpoints_data:
    doc.add_heading(modulo, level=2)
    table = doc.add_table(rows=len(endpoints)+1, cols=4)
    table.style = 'Light Grid Accent 1'
    headers = ['Metodo', 'Ruta', 'Descripcion', 'Requiere Auth']
    for j, h in enumerate(headers):
        table.cell(0, j).text = h
        for run in table.cell(0, j).paragraphs[0].runs:
            run.bold = True
    for i, (met, ruta, desc, auth) in enumerate(endpoints, 1):
        table.cell(i, 0).text = met
        table.cell(i, 1).text = ruta
        table.cell(i, 2).text = desc
        table.cell(i, 3).text = auth
    doc.add_paragraph()

doc.add_page_break()

# ===== 7. SEGURIDAD =====
doc.add_heading('7. Seguridad', level=1)
seg_items = [
    ('Autenticacion JWT', 'Uso de Flask-JWT-Extended con tokens de acceso de 24 horas de validez. '
     'Los tokens se generan con create_access_token() y se verifican con @jwt_required().'),
    ('Hash de Contrasenas', 'Las contrasenas se almacenan usando werkzeug.security.generate_password_hash() '
     '(algoritmo pbkdf2:sha256 por defecto). Nunca se almacenan en texto plano.'),
    ('Proteccion de Endpoints', '22 endpoints requieren autenticacion JWT. Los endpoints publicos son solo '
     'GET de listado y los de autenticacion (register, login).'),
    ('Validacion de Datos', 'Todos los endpoints validan campos requeridos con validate_required(). '
     'Errores devueltos como JSON con codigo HTTP 400.'),
    ('Seguridad de Base de Datos', 'Conexiones SSL obligatorias a PostgreSQL. '
     'Variables de entorno gestionadas via .env, excluidas de Git.'),
    ('Borrado en Cascada', 'Las relaciones entre tablas tienen cascade="all, delete-orphan" '
     'para evitar datos huerfanos al eliminar un usuario.'),
]
for titulo, desc in seg_items:
    p = doc.add_paragraph()
    run = p.add_run(titulo + ': ')
    run.bold = True
    p.add_run(desc)

# ===== 8. TECNOLOGIAS =====
doc.add_heading('8. Tecnologias Utilizadas', level=1)
techs = [
    ('Backend', 'Flask 3.0.3 (Python 3.10+)'),
    ('Base de Datos SQL', 'PostgreSQL con SQLAlchemy 2.0'),
    ('Base de Datos NoSQL', 'MongoDB con PyMongo 4.17'),
    ('Autenticacion', 'Flask-JWT-Extended 4.7.3'),
    ('Hash de Contrasenas', 'Werkzeug Security'),
    ('DB Driver PostgreSQL', 'Psycopg2-binary 2.9.12'),
    ('Variables de Entorno', 'Python-dotenv 1.0.1'),
    ('Serializacion', 'JSON nativo de Flask'),
    ('Control de Versiones', 'Git + GitHub'),
    ('Entorno de Desarrollo', 'Visual Studio Code'),
]
table = doc.add_table(rows=len(techs)+1, cols=2)
table.style = 'Light Grid Accent 1'
table.cell(0, 0).text = 'Componente'
table.cell(0, 1).text = 'Tecnologia'
for run in table.cell(0, 0).paragraphs[0].runs:
    run.bold = True
for run in table.cell(0, 1).paragraphs[0].runs:
    run.bold = True
for i, (comp, tech) in enumerate(techs, 1):
    table.cell(i, 0).text = comp
    table.cell(i, 1).text = tech

# ===== 9. DEPENDENCIAS =====
doc.add_heading('9. Dependencias (requirements.txt)', level=1)
doc.add_paragraph('Lista completa de paquetes Python necesarios para ejecutar el proyecto:')
deps = [
    ('Flask', '3.0.3', 'Framework web'),
    ('Flask-SQLAlchemy', '3.1.1', 'ORM para PostgreSQL'),
    ('psycopg2-binary', '2.9.12', 'Driver PostgreSQL'),
    ('pymongo', '4.17.0', 'Cliente MongoDB'),
    ('python-dotenv', '1.0.1', 'Carga de variables de entorno'),
    ('Flask-JWT-Extended', '4.7.3', 'Autenticacion JWT'),
]
table = doc.add_table(rows=len(deps)+1, cols=3)
table.style = 'Light Grid Accent 1'
for j, h in enumerate(['Paquete', 'Version', 'Proposito']):
    table.cell(0, j).text = h
    for run in table.cell(0, j).paragraphs[0].runs:
        run.bold = True
for i, (pkg, ver, prop) in enumerate(deps, 1):
    table.cell(i, 0).text = pkg
    table.cell(i, 1).text = ver
    table.cell(i, 2).text = prop

# FOOTER
doc.add_paragraph()
doc.add_paragraph()
footer = doc.add_paragraph()
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = footer.add_run('--- Fin del documento ---')
run.font.color.rgb = RGBColor(150, 150, 150)
run.italic = True

doc.save('informe_requisitos.docx')
print('DOCX generado: informe_requisitos.docx')
